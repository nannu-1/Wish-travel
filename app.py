import gradio as gr
from main import app  # Ensure app is imported correctly
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OpenAIEmbeddings

import gradio as gr
from datetime import datetime
from main import app
from docx import Document
from fpdf import FPDF

def save_itinerary_file(city, country, itinerary, file_format):
    filename_base = f"{city}_{country}".replace(" ", "_")

    if file_format == "TXT":
        txt_file = f"{filename_base}_itinerary.txt"
        with open(txt_file, "w", encoding="utf-8") as f:
            f.write(itinerary)
        return txt_file

    elif file_format == "DOCX":
        doc = Document()
        doc.add_heading("WanderMind Itinerary", 0)
        doc.add_paragraph(itinerary)
        docx_file = f"{filename_base}_itinerary.docx"
        doc.save(docx_file)
        return docx_file

    elif file_format == "PDF":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, itinerary)
        pdf_file = f"{filename_base}_itinerary.pdf"
        pdf.output(pdf_file)
        return pdf_file

def run_planner(city, country, interests, travel_date_text, budget, file_format):
    try:
        # Parse date in DD-MM-YYYY format
        travel_date = datetime.strptime(travel_date_text.strip(), "%d-%m-%Y").date()
        today = datetime.today().date()
        if travel_date < today:
            return "⚠️ Travel date must be today or in the future.", None
    except ValueError:
        return "❌ Invalid date format. Please use DD-MM-YYYY.", None

    # Format date back to YYYY-MM-DD for processing
    formatted_date = travel_date.strftime("%Y-%m-%d")

    input_data = {
        "messages": [],
        "city": city,
        "country": country,
        "interests": [i.strip() for i in interests.split(",")],
        "travel_dates": formatted_date,
        "budget": budget,
        "itinerary": ""
    }

    result = app.invoke(input_data)
    itinerary = result["itinerary"]

    file_path = save_itinerary_file(city, country, itinerary, file_format)
    return itinerary, file_path

ui = gr.Interface(
    fn=run_planner,
    inputs=[
        gr.Textbox(label="City"),
        gr.Textbox(label="Country"),
        gr.Textbox(label="Interests (comma-separated)"),
        gr.Textbox(label="Travel Date (DD-MM-YYYY)"),  # 👈 Updated format
        gr.Dropdown(["low", "medium", "high"], label="Budget"),
        gr.Radio(["TXT", "DOCX", "PDF"], label="Download Format"),
    ],
    outputs=[
        gr.Textbox(label="Generated Itinerary"),
        gr.File(label="Download File"),
    ],
    title="🌍 WanderMind: AI Travel Planner",
    description="Enter a future travel date in DD-MM-YYYY format. WanderMind will create a custom itinerary for you.",
)

ui.launch(share=True)


